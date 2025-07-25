import streamlit as st
from pdf2image import convert_from_path
import pytesseract
from PIL import Image
from transformers import pipeline
from docx import Document
import tempfile
import os
from transformers import pipeline

summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6")


# Functions
def convert_all_to_images(file_path):
    image_list = []
    if file_path.lower().endswith(".pdf"):
        images = convert_from_path(file_path, dpi=500)
        image_list.extend(images)
    elif file_path.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp')):
        image = Image.open(file_path).convert("RGB")
        image_list.append(image)
    return image_list

def apply_ocr(images):
    extracted_text = ""
    for img in images:
        text = pytesseract.image_to_string(img, lang="eng")
        extracted_text += text + "\n"
    return extracted_text.strip()

def preprocess_text(text):
    return ' '.join(text.replace('\n', ' ').split())

def summarize_text(text):
    max_chunk = 1024
    sentences = text.split('. ')
    chunks = []
    current_chunk = ''
    for sentence in sentences:
        if len(current_chunk) + len(sentence) <= max_chunk:
            current_chunk += sentence + '. '
        else:
            chunks.append(current_chunk)
            current_chunk = sentence + '. '
    chunks.append(current_chunk)
    summary = ''
    for chunk in chunks:
        summary += summarizer(chunk, max_length=150, min_length=40, do_sample=False)[0]['summary_text'] + ' '
    return summary.strip()

def save_outputs(text, summary):
    with open("extracted_text.txt", "w") as f:
        f.write(text)
    doc = Document()
    doc.add_heading("Summary", 0)
    doc.add_paragraph(summary)
    doc.save("summary.docx")
    return "extracted_text.txt", "summary.docx"

# Streamlit UI
st.set_page_config(page_title="Handwritten Text Summarizer", layout="centered")

st.title("📝 Handwritten PDF/Image Summarizer")
st.write("Upload a **handwritten PDF or image**, and this app will extract the text and provide a clean **summary**.")

uploaded_file = st.file_uploader("Upload PDF/Image", type=["pdf", "png", "jpg", "jpeg", "bmp"])

if uploaded_file is not None:
    with tempfile.NamedTemporaryFile(delete=False) as tf:
        tf.write(uploaded_file.getvalue())
        path = tf.name

    with st.spinner("Processing..."):
        images = convert_all_to_images(path)
        extracted = apply_ocr(images)
        cleaned = preprocess_text(extracted)
        summary = summarize_text(cleaned)
        txt_path, docx_path = save_outputs(cleaned, summary)

    st.success("✅ Text and Summary Ready!")

    st.subheader("🧾 Extracted Text")
    st.text_area("Text Output", cleaned, height=300)

    st.subheader("🧠 Summary")
    st.text_area("Summary Output", summary, height=200)

    st.download_button("📄 Download Text (.txt)", open(txt_path, "rb"), file_name="extracted_text.txt")
    st.download_button("📄 Download Summary (.docx)", open(docx_path, "rb"), file_name="summary.docx")
