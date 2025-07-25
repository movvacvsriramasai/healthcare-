# app.py

import streamlit as st
import os
import openai
from pdf2image import convert_from_path
from PIL import Image
from docx import Document
import base64
import tempfile

# --- Set your API key here (or use secrets for Streamlit Cloud) ---
openai.api_key = st.secrets["OPENAI_API_KEY"] if "OPENAI_API_KEY" in st.secrets else os.getenv("OPENAI_API_KEY")

# --- Functions ---
def save_uploaded_file(uploaded_file):
    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp_file:
        tmp_file.write(uploaded_file.getbuffer())
        return tmp_file.name

def convert_pdf_to_images(pdf_path):
    images = convert_from_path(pdf_path, dpi=300)
    image_paths = []
    for i, img in enumerate(images):
        path = f"page_{i+1}.png"
        img.save(path, "PNG")
        image_paths.append(path)
    return image_paths

def encode_image(image_path):
    with open(image_path, "rb") as img_file:
        return base64.b64encode(img_file.read()).decode("utf-8")

def extract_text_with_gpt_vision(image_path):
    base64_image = encode_image(image_path)
    try:
        response = openai.chat.completions.create(
            model="gpt-4-vision-preview",
            messages=[
                {"role": "system", "content": "You extract clean digital text from handwritten images."},
                {"role": "user", "content": [
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{base64_image}"}}
                ]}
            ],
            max_tokens=2000
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Error extracting text: {str(e)}"

def summarize_text_with_gpt(text):
    try:
        response = openai.chat.completions.create(
            model="gpt-4-turbo",
            messages=[
                {"role": "system", "content": "You summarize long text clearly and concisely."},
                {"role": "user", "content": f"Summarize this content:\n\n{text}"}
            ],
            max_tokens=1024
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Error generating summary: {str(e)}"

def generate_docx(text, summary):
    doc = Document()
    doc.add_heading("Extracted Text", level=1)
    doc.add_paragraph(text)
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary)
    file_path = "output.docx"
    doc.save(file_path)
    return file_path

# --- Streamlit UI ---
st.title("📝 Handwritten OCR + Summarizer (OpenAI)")

uploaded_files = st.file_uploader("Upload a PDF or image", type=["pdf", "png", "jpg", "jpeg"], accept_multiple_files=True)

if uploaded_files:
    full_text = ""

    for file in uploaded_files:
        st.write(f"Processing: {file.name}")
        file_path = save_uploaded_file(file)

        image_paths = []
        if file.name.endswith(".pdf"):
            image_paths = convert_pdf_to_images(file_path)
        else:
            image_paths = [file_path]

        for img_path in image_paths:
            st.image(img_path, caption="Processed Image", use_column_width=True)
            extracted_text = extract_text_with_gpt_vision(img_path)
            st.text_area("Extracted Text", extracted_text, height=200)
            full_text += extracted_text + "\n\n"

    st.subheader("📄 Generating Summary...")
    summary = summarize_text_with_gpt(full_text)
    st.text_area("Summary", summary, height=200)

    # Download links
    with open("output.txt", "w") as f:
        f.write(full_text)

    with open("summary.txt", "w") as f:
        f.write(summary)

    docx_path = generate_docx(full_text, summary)

    st.download_button("Download Extracted Text (.txt)", data=open("output.txt", "rb"), file_name="output.txt")
    st.download_button("Download Summary (.txt)", data=open("summary.txt", "rb"), file_name="summary.txt")
    st.download_button("Download Full Document (.docx)", data=open(docx_path, "rb"), file_name="output.docx")


#Requirements

streamlit
openai
pdf2image
pillow
python-docx
poppler-utils
torch
